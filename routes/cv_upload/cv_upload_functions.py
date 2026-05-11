import io
import re
import easyocr
import pdfplumber

from typing import Optional

from functions.logger import logger

_easyocr_reader = None


def _get_easyocr_reader():
    global _easyocr_reader
    if _easyocr_reader is None:
        logger("CV_UPLOAD", "Initializing EasyOCR reader (first-time load)", level="INFO")
        _easyocr_reader = easyocr.Reader(["en"], gpu=False, verbose=False)
    return _easyocr_reader


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF using both pdfplumber and EasyOCR, then return whichever
    yields more characters. pdfplumber handles selectable text (fast, preserves order);
    EasyOCR captures text embedded in images or complex layouts that pdfplumber misses.
    """
    # --- pdfplumber ---
    plumber_text = ""
    try:
        logger("CV_UPLOAD", "Extracting PDF text with pdfplumber", level="DEBUG")
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        plumber_text = "\n".join(pages).strip()
        logger("CV_UPLOAD", f"pdfplumber extracted {len(plumber_text)} chars", level="DEBUG")
    except Exception as e:
        logger("CV_UPLOAD", f"pdfplumber failed: {e}", level="DEBUG")

    # --- EasyOCR (always runs) ---
    ocr_text = ""
    try:
        import numpy as np
        from pdf2image import convert_from_bytes
        logger("CV_UPLOAD", "Converting PDF pages to images for EasyOCR", level="DEBUG")
        images = convert_from_bytes(file_bytes, dpi=200)
        reader = _get_easyocr_reader()
        page_texts = []
        for i, img in enumerate(images):
            blocks = reader.readtext(np.array(img), detail=0, paragraph=True)
            page_texts.append("\n".join(blocks))
            logger("CV_UPLOAD", f"EasyOCR page {i + 1}: {len(blocks)} text blocks", level="DEBUG")
        ocr_text = "\n".join(page_texts).strip()
        logger("CV_UPLOAD", f"EasyOCR extracted {len(ocr_text)} chars", level="DEBUG")
    except Exception as e:
        logger("CV_UPLOAD", f"EasyOCR failed: {e}", level="DEBUG")

    # Return whichever extraction is richer
    if ocr_text and len(ocr_text) >= len(plumber_text):
        logger("CV_UPLOAD", "Using EasyOCR result (richer or equal)", level="DEBUG")
        return ocr_text
    if plumber_text:
        logger("CV_UPLOAD", "Using pdfplumber result (richer)", level="DEBUG")
        return plumber_text
    raise RuntimeError("Failed to extract text from PDF (both pdfplumber and EasyOCR failed)")


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX including table cells (skills are often in tables)."""
    try:
        import docx
        logger("CV_UPLOAD", "Extracting DOCX text with python-docx", level="DEBUG")
        doc = docx.Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        text = "\n".join(parts).strip()
        logger("CV_UPLOAD", f"DOCX extracted {len(text)} chars", level="DEBUG")
        return text
    except Exception as e:
        logger("CV_UPLOAD", f"DOCX extraction failed: {e}", level="ERROR")
        raise RuntimeError(f"Failed to extract text from DOCX: {e}")


def _extract_text_from_image(file_bytes: bytes) -> str:
    """OCR chain for image file uploads (PNG/JPG/etc): EasyOCR → Tesseract → Gemini."""
    logger("CV_UPLOAD", "Starting image OCR extraction", level="DEBUG")

    try:
        import numpy as np
        from PIL import Image
        logger("CV_UPLOAD", "Attempting EasyOCR on image", level="DEBUG")
        image = Image.open(io.BytesIO(file_bytes))
        if image.mode != "RGB":
            image = image.convert("RGB")
        reader = _get_easyocr_reader()
        blocks = reader.readtext(np.array(image), detail=0, paragraph=True)
        text = "\n".join(blocks).strip()
        if text:
            logger("CV_UPLOAD", f"EasyOCR extracted {len(text)} chars", level="DEBUG")
            return text
        logger("CV_UPLOAD", "EasyOCR returned empty text", level="DEBUG")
    except Exception as e:
        logger("CV_UPLOAD", f"EasyOCR failed: {e}", level="DEBUG")

    try:
        from PIL import Image
        import pytesseract
        logger("CV_UPLOAD", "Attempting Tesseract OCR", level="DEBUG")
        image = Image.open(io.BytesIO(file_bytes))
        if image.mode != "RGB":
            image = image.convert("RGB")
        text = pytesseract.image_to_string(image).strip()
        if text:
            logger("CV_UPLOAD", f"Tesseract extracted {len(text)} chars", level="DEBUG")
            return text
        logger("CV_UPLOAD", "Tesseract returned empty text", level="DEBUG")
    except Exception as e:
        logger("CV_UPLOAD", f"Tesseract OCR failed: {e}", level="DEBUG")

    try:
        import os
        from google import genai
        from PIL import Image
        logger("CV_UPLOAD", "Attempting Gemini OCR", level="DEBUG")
        project_id = os.getenv("GOOGLE_PROJECT_ID")
        if project_id:
            client = genai.Client(
                vertexai=True,
                project=project_id,
                location=os.getenv("GOOGLE_LOCATION", "us-central1"),
            )
        else:
            api_key = os.getenv("GEMINI_API_KEY")
            if not api_key:
                raise RuntimeError("No Gemini credentials configured")
            client = genai.Client(api_key=api_key)
        image = Image.open(io.BytesIO(file_bytes))
        response = client.models.generate_content(
            model=os.getenv("GOOGLE_LLM", "gemini-2.5-flash"),
            contents=["Extract all visible text from this image. Return only the extracted text.", image],
            config={"temperature": 0.0, "max_output_tokens": 2048},
        )
        text = response.text.strip()
        if text:
            logger("CV_UPLOAD", f"Gemini OCR extracted {len(text)} chars", level="DEBUG")
            return text
        logger("CV_UPLOAD", "Gemini OCR returned empty text", level="DEBUG")
    except Exception as e:
        logger("CV_UPLOAD", f"Gemini OCR failed: {e}", level="DEBUG")

    raise RuntimeError("All OCR methods failed (EasyOCR, Tesseract, Gemini).")
